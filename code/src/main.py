import os
import re
import datetime
import logging
import torch
import json
import fitz  # PyMuPDF for PDF text extraction
import pdfplumber  # For extracting tables
import docx
from email import policy
from email.parser import BytesParser
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from pathlib import Path
from difflib import SequenceMatcher


# Define mappings for Request Type and Sub-Request Type
REQUEST_TYPE_MAP = {
    "Adjustment": ["—"],
    "AU Transfer": ["—"],
    "Closing Notice": ["Reallocation Fees", "Amendment Fees", "Reallocation Principal"],
    "Commitment Change": ["Cashless Roll", "Decrease", "Increase"],
    "Fee Payment": ["Ongoing Fee", "Letter of Credit Fee"],
    "Money Movement - Inbound": ["Principal", "Interest", "Principal + Interest", "Principal + Interest + Fee"],
    "Money Movement - Outbound": ["Timebound", "Foreign Currency"],
}

# Define mappings for fields corresponding to each of the Request Type and Sub-Request Type
EXPECTED_FIELDS_MAP = {
    "Adjustment": {
        "—": ["bank_name", "effective_date", "adjustment_amount", "commitment_share", "new_commitment_share"]
    },
    "AU Transfer": {
        "—": ["bank_name", "transfer_amount", "transfer_date", "reference"]
    },
    "Closing Notice": {
        "Reallocation Fees": ["bank_name", "deal_cusip", "facility_cusip", "sofr_payment", "fee_amount"],
        "Amendment Fees": ["bank_name", "deal_isin", "facility_isin", "lender_mei", "amendment_fee"],
        "Reallocation Principal": ["bank_name", "deal_cusip", "facility_cusip", "principal_amount", "effective_date"]
    },
    "Commitment Change": {
        "Cashless Roll": ["bank_name", "effective_date", "global_principal_balance", "loan_repricing_date"],
        "Decrease": ["bank_name", "facility_cusip", "lender_mei", "decrease_amount", "previous_commitment_share"],
        "Increase": ["bank_name", "facility_cusip", "lender_mei", "increase_amount", "new_commitment_share"]
    },
    "Fee Payment": {
        "Ongoing Fee": ["bank_name", "fee_amount", "fee_type", "payment_date", "reference"],
        "Letter of Credit Fee": ["bank_name", "fee_amount", "letter_of_credit_id", "payment_date", "bank_name"]
    },
    "Money Movement - Inbound": {
        "Principal": ["global_principal_balance", "lender_share_principal_balance", "principal_payment", "effective_date"],
        "Interest": ["sofr_payment", "interest_payment", "repayment_date", "borrower_share"],
        "Principal + Interest": ["principal_payment", "interest_payment", "repayment_date"],
        "Principal + Interest + Fee": ["principal_payment", "interest_payment", "fee_amount", "repayment_date"]
    },
    "Money Movement - Outbound": {
        "Timebound": ["bank_name", "transfer_amount", "transfer_date", "maturity_date", "account_name"],
        "Foreign Currency": ["bank_name", "transfer_amount", "currency_type", "exchange_rate", "account_number"]
    }
}


# ✅ Logging Function
def log_and_print(message, level="info"):
    timestamp = datetime.datetime.now().strftime("[%Y-%m-%d %H:%M:%S]")
    log_message = f"{timestamp} {message}"

    if level == "info":
        logging.info(log_message)
    elif level == "warning":
        logging.warning(log_message)
    elif level == "error":
        logging.error(log_message)

    print(log_message)


# ✅ Setup Logging
logging.basicConfig(filename="processing.log", level=logging.INFO, format="%(asctime)s - %(message)s")


# ✅ Load Mistral-7B with 4-bit Quantization (Optimized for 4GB GPU)
# Define Constants
DATA_FOLDER = "Input"
MODEL_NAME = "mistralai/Mistral-7B-Instruct-v0.1"

quantization_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_compute_dtype=torch.float16,
    bnb_4bit_use_double_quant=True
)

tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForCausalLM.from_pretrained(
    MODEL_NAME,
    device_map="auto",
    quantization_config=quantization_config
).to("cuda")


# ✅ Extract Text Functions (PDF, DOCX, EML)
def extract_text_pdf(file_path):
    """Extracts text from a PDF, including tables."""
    try:
        text = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text.append(page.get_text())

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                table = page.extract_table()
                if table:
                    text.append("\n".join([" | ".join(row) for row in table if row]))

        return "\n".join(text).strip()
    except Exception as e:
        log_and_print(f"❌ Error extracting PDF: {e}", "error")
        return ""


def extract_text_docx(file_path):
    """Extracts text from a DOCX file, including tables."""
    try:
        doc = docx.Document(file_path)
        text = [para.text for para in doc.paragraphs]

        for table in doc.tables:
            for row in table.rows:
                text.append(" | ".join([cell.text.strip() for cell in row.cells]))

        return "\n".join(text).strip()
    except Exception as e:
        log_and_print(f"❌ Error extracting DOCX: {e}", "error")
        return ""


def extract_text_eml(file_path):
    """Extracts text from EML files and removes duplicate content."""
    try:
        with open(file_path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)

        email_body = msg.get_body(preferencelist=('plain')).get_content().strip() if msg.get_body() else ""
        extracted_texts = []
        unique_texts = set()

        if email_body:
            unique_texts.add(email_body)
            extracted_texts.append({"source": "email_body", "text": email_body})

        for part in msg.iter_attachments():
            filename = part.get_filename()
            content_type = part.get_content_type()

            if filename and content_type:
                attachment_path = os.path.join("temp_attachments", filename)
                os.makedirs("temp_attachments", exist_ok=True)

                with open(attachment_path, "wb") as f:
                    f.write(part.get_payload(decode=True))

                extracted_text = extract_text_pdf(attachment_path) if filename.endswith(".pdf") else extract_text_docx(attachment_path)

                if extracted_text and extracted_text not in unique_texts:
                    unique_texts.add(extracted_text)
                    extracted_texts.append({"source": filename, "text": extracted_text})

        return extracted_texts
    except Exception as e:
        log_and_print(f"❌ Error extracting EML: {e}", "error")
        return []

# 🔹 Extract Text Based on File Type
def extract_text(file_path):
    """Extracts text from PDFs, DOCX, and EML files while handling duplicates."""
    file_handlers = {
        ".pdf": extract_text_pdf,
        ".docx": extract_text_docx,
        ".eml": extract_text_eml
    }

    file_extension = Path(file_path).suffix.lower()
    extract_function = file_handlers.get(file_extension)

    if extract_function:
        extracted_text = extract_function(file_path)
    else:
        log_and_print(f"⚠️ Unsupported file format: {file_extension}", "warning")
        return []

    # Ensure extracted_text is always a list of dictionaries
    if isinstance(extracted_text, str):  
        extracted_text = [{"source": file_path, "text": extracted_text}]
    
    return extracted_text


def extract_financial_entities(text):
    """Extracts financial key-value pairs from raw text using Mistral-7B."""
    try:
        prompt = f"""
        Below is a financial document. Extract key-value pairs and return output in valid JSON format.
        Ensure that the JSON follows the correct structure and contains meaningful financial entities.

        Financial Document:
        {text}

        JSON Output:
        """

        # Tokenize input
        inputs = tokenizer(prompt, return_tensors="pt").to("cuda")

        # Generate response
        with torch.no_grad():
            outputs = model.generate(**inputs, max_new_tokens=1200)

        # Decode response
        extracted_data = tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

        # ✅ Extract only JSON using regex
        match = re.search(r'\{.*\}', extracted_data, re.DOTALL)
        if match:
            json_string = match.group(0)
        else:
            return {"error": "No valid JSON found", "raw_output": extracted_data}

        # ✅ Convert to structured JSON
        try:
            structured_data = json.loads(json_string)
        except json.JSONDecodeError:
            return {"error": "Failed to parse extracted JSON", "raw_output": json_string}

        return structured_data

    except Exception as e:
        return {"error": str(e)}

def flatten_dict(d, parent_key=""):
    """Flatten nested dictionary keys using underscore notation."""
    items = {}
    for k, v in d.items():
        new_key = f"{parent_key}_{k.lower().replace(' ', '_')}".strip("_")  # Normalize field names
        if isinstance(v, dict):
            items.update(flatten_dict(v, new_key))  # Recursive flattening
        else:
            items[new_key] = v  # Store non-dictionary values
    return items

def count_matching_fields(extracted_data, expected_fields):
    """Recursively count matching fields in extracted data (handles nested objects)."""
    flat_extracted_data = flatten_dict(extracted_data)  # Flatten nested fields

    # Normalize expected field names (convert spaces to underscores)
    expected_fields_normalized = {field.lower().replace(" ", "_") for field in expected_fields}

    # Count matches based on normalized field names
    matched_fields = sum(1 for field in expected_fields_normalized if field in flat_extracted_data)

    total_fields = len(expected_fields_normalized)
    
    if total_fields == 0:
        return 0.0  # Avoid division by zero

    confidence_score = round((matched_fields / total_fields) * 100, 2)
    
    return confidence_score



# ✅ Extract & Classify Request Type
def classify_request_type_based_on_fields(extracted_data):
    """Determine the most likely Request Type and Sub-Request Type based on extracted fields."""
    
    # 🔍 Normalize extracted data first
    normalized_data = {key: normalize_text(str(value)) for key, value in flatten_dict(extracted_data).items()}
    #print(f"\n🔍 Normalized Extracted Data for Classification:\n{json.dumps(normalized_data, indent=4)}")

    best_match = None
    best_confidence = 0.0

    for request_type, sub_types in EXPECTED_FIELDS_MAP.items():
        for sub_request_type, expected_fields in sub_types.items():
            matched_fields = sum(1 for field in expected_fields if field in normalized_data and normalized_data[field])
            total_fields = len(expected_fields)
            confidence_score = round((matched_fields / max(total_fields, 1)) * 100, 2)

            if confidence_score > best_confidence:
                best_match = (request_type, sub_request_type)
                best_confidence = confidence_score

    return {
        "Request Type": best_match[0] if best_match else "-",
        "Sub Request Type": best_match[1] if best_match else "-",
        "Confidence Score": best_confidence / 100  # Normalize to 0-1 scale
    }



# ✅ Update JSON Format & Classify
def normalize_text(text):
    """Normalize text for duplicate detection while preserving financial symbols."""
    text = text.replace("\u2014", "-")  # Convert em-dash to normal dash
    text = text.lower().strip()  # Convert to lowercase & remove leading/trailing spaces
    text = re.sub(r"^[-]+|[-]+$", "", text)  # Remove leading/trailing separators (like "---")
    text = re.sub(r"\s+", " ", text)  # Collapse multiple spaces into one
    text = re.sub(r"[^\w\s$%#@.,:/-]", "", text)  # Keep financial symbols & email-friendly chars
    return text

def is_duplicate(existing_entries, new_entry):
    """Check if new_entry is a duplicate based on extracted field similarity and content matching."""
    
    new_extracted_text = new_entry.get("Extracted Text", {})
    new_fields_count = len(flatten_dict(new_extracted_text))

    print(f"\n🔍 Checking duplicate for: {new_entry['source']} ({new_fields_count} fields)")

    best_match = None
    best_similarity = 0.0

    for existing_entry in existing_entries:
        if existing_entry is new_entry:
            continue  # ✅ Skip self-comparison

        existing_extracted_text = existing_entry.get("Extracted Text", {})
        existing_fields_count = len(flatten_dict(existing_extracted_text))

        # ✅ **Compute Similarity on Extracted Fields**
        similarity = SequenceMatcher(
            None, json.dumps(existing_extracted_text, sort_keys=True),
            json.dumps(new_extracted_text, sort_keys=True)
        ).ratio()

        print(f"   🔹 Comparing with: {existing_entry['source']} ({existing_fields_count} fields) ➡️ Similarity: {similarity:.2f}")

        # ✅ Track the best matching duplicate
        if similarity > best_similarity:
            best_match = existing_entry
            best_similarity = similarity

    # ✅ **Mark duplicate if similarity > 0.6**
    if best_similarity > 0.6:
        print(f"   ✅ Duplicate Found: {new_entry['source']} vs {best_match['source']} (Similarity: {best_similarity:.2f})")

        best_match_fields = len(flatten_dict(best_match.get("Extracted Text", {})))

        if new_fields_count > best_match_fields:
            best_match["Duplicate Content"] = "Yes"
            print(f"   ➡️ Marking EXISTING entry {best_match['source']} as duplicate.")
            return False  # ✅ Keep the new entry
        else:
            print(f"   ➡️ Marking NEW entry {new_entry['source']} as duplicate.")
            return True  # ✅ Mark new entry as duplicate

    print(f"✅ No duplicates found for {new_entry['source']}.")
    return False  # ✅ Not a duplicate

def update_extracted_json(extracted_data):
    """Classify request type, compute confidence score, and detect duplicates properly."""
    final_output = {
        "Raw Data": [],
        "Extracted Data": []
    }

    unique_entries = []  # ✅ Track unique entries

    print("\n🔍 Starting Data Extraction & Classification...\n")

    for doc in extracted_data:
        text_content = doc.get("text", "").strip()
        source = doc.get("source", "unknown")

        if not text_content:
            print(f"⚠️ Skipping empty document: {source}")
            continue  

        print(f"\n🔍 Processing Source: {source}")

        # ✅ Extract financial entities
        extracted_financial_data = extract_financial_entities(text_content)

        print(f"🔹 Extracted Data:\n{json.dumps(extracted_financial_data, indent=4)}")

        if isinstance(extracted_financial_data, dict) and extracted_financial_data:
            classification_result = classify_request_type_based_on_fields(extracted_financial_data)
        else:
            classification_result = {
                "Request Type": "-",
                "Sub Request Type": "-",
                "Confidence Score": 0.0
            }

        print(f"🔍 Classification Result: {classification_result}")

        # ✅ Replace `\u2014` (em dash) with `-` (hyphen)
        classification_result["Request Type"] = classification_result["Request Type"].replace("\u2014", "-")
        classification_result["Sub Request Type"] = classification_result["Sub Request Type"].replace("\u2014", "-")

        new_entry = {
            "source": source,
            "text": text_content,
            "Extracted Text": extracted_financial_data,
            "Request Type": classification_result["Request Type"],
            "Sub Request Type": classification_result["Sub Request Type"],
            "Confidence Score": classification_result["Confidence Score"],
            "Duplicate Content": "No"  
        }

        unique_entries.append(new_entry)

    print("\n🔎 Checking for duplicates after extraction...\n")

    for entry in unique_entries:
        is_duplicate_flag = "Yes" if is_duplicate(unique_entries, entry) else "No"
        entry["Duplicate Content"] = is_duplicate_flag
        print(f"✅ Final Decision for {entry['source']}: Duplicate = {is_duplicate_flag}")

        final_output["Raw Data"].append({
            "source": entry["source"],
            "text": entry["text"],
            "Duplicate Content": is_duplicate_flag
        })

        if is_duplicate_flag == "No":
            final_output["Extracted Data"].append({
                "source": entry["source"],
                "Extracted Text": entry["Extracted Text"],
                "Request Type": entry["Request Type"],
                "Sub Request Type": entry["Sub Request Type"],
                "Confidence Score": entry["Confidence Score"]
            })

    print("\n✅ Final Processed JSON:")
    print(json.dumps(final_output, indent=4))

    return final_output

# 🔹 Save JSON Output
def save_output_as_json(data, filename):
    """Saves extracted key-value pairs as JSON with file extension prefix and timestamp."""
    output_folder = "Output"
    os.makedirs(output_folder, exist_ok=True)

    file_extension = Path(filename).suffix.lstrip(".")  # Extract file extension (e.g., "pdf")
    base_filename = Path(filename).stem  # Extract filename without extension

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")  # Timestamp format
    json_filename = f"{file_extension}_{base_filename}_{timestamp}.json"  # Example: pdf_sample_20240325_153045.json

    json_path = os.path.join(output_folder, json_filename)

    # Save JSON data
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(data, json_file, indent=4)

    log_and_print(f"📂 JSON saved: {json_path}")



# 🔹 Main Execution
def main():
    for file_name in os.listdir(DATA_FOLDER):
        extracted_text = extract_text(os.path.join(DATA_FOLDER, file_name))
        if extracted_text:
            save_output_as_json(update_extracted_json(extracted_text), file_name)


if __name__ == "__main__":
    main()

